import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import os
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import json
import sys
import logging


# 配置logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# 创建控制台处理器
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.DEBUG)

# 创建格式器
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)

# 添加处理器到logger
logger.addHandler(console_handler)


# 定义重试策略
retry = Retry(
    total=3,  # 总重试次数（含首次共4次）
    backoff_factor=1,  # 指数退避：0s,1s,2s,4s...
    status_forcelist=[500,502,503,504,429],  # 这些状态码触发重试
    allowed_methods=["POST"],  # 允许重试 POST（谨慎！）
)


BASE_PROTOCOL = "http"
BASE_HOST_DOMAIN = "117.50.220.141"
# BASE_HOST_DOMAIN = "106.75.46.58"
BASE_HOST_PORT = ":15000"
# BASE_HOST_PORT = ":18085"
TRANSLATION_API = f"{BASE_PROTOCOL}://{BASE_HOST_DOMAIN}{BASE_HOST_PORT}/translate_batch_v2"


def get_translated_result(
        request_url=TRANSLATION_API,
            content=None,
            language_direction=("en", "zh")):

    headers = {"Content-Type": "application/json"}

    body_data = {
        "domain": "medical",
        "qs": [content],
        "source": language_direction[0],
        "target": language_direction[1],}

    params = {
        "instance_id": "xxx",
        "application_id": "xxx",}

    adapter = HTTPAdapter(max_retries=retry)
    session = requests.Session()
    session.mount("https://", adapter)
    session.mount("http://", adapter)

    try:
        resp  = session.post(
            request_url,
            headers=headers,
            json=body_data,
            params=params,
            timeout=30)
        
        if resp.status_code == 200:
            translation_result = resp.json()['translation'][0]
            logger.debug(f"翻译结果: {translation_result}")
            return translation_result
        else:
            logger.error(f"请求失败! 状态码: {resp.status_code}")
            logger.error(f"错误信息: {resp.text}")
            return ""
    except requests.exceptions.RequestException as e:
        logger.error(f"请求异常: {e}")
        return ""


def extract_docx_content(file_path):
    """
    提取 .docx 文件中的所有文本内容
    """
    try:
        # 打开文档
        document = docx.Document(file_path)

        all_content = []

        # 提取段落内容
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                all_content.append(paragraph.text.strip())

        # 提取表格内容
        for element in document.element.body:
            if isinstance(element, CT_Tbl):
                table = Table(element, document)
                for row in table.rows:
                    for cell in row.cells:
                        cell_content = extract_paragraphs_from_cell(cell)
                        if cell_content.strip():
                            all_content.append(cell_content.strip())

        return all_content

    except Exception as e:
        logger.error(f"读取文档失败: {e}")
        return []

def extract_paragraphs_from_cell(cell):
    """
    从表格单元格中提取文本内容
    """
    content_parts = []
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            content_parts.append(paragraph.text.strip())
    return " ".join(content_parts)


def extract_docx_content_simple(file_path):
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return paragraphs
    except Exception as e:
        logger.error(f"读取文档失败: {e}")
        return []


def save_content_to_file(content_list, output_filename, output_translated_filename):
    """
    将提取的内容保存到文件
    """
    with open(output_filename, 'w', encoding='utf-8') as output_file, \
        open(output_translated_filename, 'w', encoding='utf-8') as output_translated_file:
        for content in content_list:
            translated_content = get_translated_result(content=content)
            output_file.write(content + '\n')
            output_translated_file.write(translated_content + '\n')


def main():
    # 需要安装 python-docx 库: pip install python-docx
    docx_files_path = [
        r"c:\Users\Tsai\Desktop\参考语料文件\1\en_extract\_q0lra3Yi2Bk-JQ3DnREA.docx",
        r"C:\Users\Tsai\Desktop\参考语料文件\1\cn_extract\sea5mk0fhEEqxJ1StuJOL.docx",
        # r"C:\Users\Tsai\Desktop\参考语料文件\2\en_extract\2-nMYCPHKNSzkJCNV7A2P.docx",
        # # r"C:\Users\Tsai\Desktop\参考语料文件\2\cn_extract\FjGdM5175vVADx9_29WpN.docx",
    ]

    for file_path in docx_files_path:
        if not os.path.exists(file_path):
            print(f"文件不存在: {file_path}")
            continue

        logger.info(f"\n处理文档: {file_path}")

        # 提取文档内容
        results = extract_docx_content(file_path)

        # 显示部分内容预览
        logger.info(f"共提取到 {len(results)} 段内容")
        # for i, content in enumerate(results[:5], 1):  # 显示前5个内容
        #     print(f"{i}. {content[:100]}{'...' if len(content) > 100 else ''}")

        # 保存到新文件
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_filename = f'extracted_{base_name}_content.txt'
        output_translated_filename = f'extracted_{base_name}_translated_content.txt'
        save_content_to_file(results, output_filename, output_translated_filename)

        logger.info(f"\n已将所有提取的内容保存到 '{output_filename}' 文件中")


if __name__ == "__main__":
    main()
