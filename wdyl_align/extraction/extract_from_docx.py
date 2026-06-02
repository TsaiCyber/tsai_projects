# coding: utf-8

import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

import os
from unicodedata import digit
import json
import sys
import re
import unicodedata

from constants import EN_SENTENCE_IGNORE_LENGTH, ZH_SENTENCE_IGNORE_LENGTH
from wdyl_logger.wdyl_logger import logger


# 正则表达式
digit_only_pattern = re.compile(
    r"^[\d\s\+\-\*\/\=\!\@\#\$\%\^\&\*\(\)\[\]\{\}\;\:\'\"\,\.\<\>\/\?\\\|`~]+$")
zh_digit_english_only_pattern = re.compile(
    r"^[A-Za-z0-9\s\.\,\!\?\;\:\'\"\(\)\[\]\{\}\<\>\@\#\$\%\^\&\*\+\-\=\/\\\|`~：（）–]+$")


def extract_paragraphs_from_cell(cell):
    """
    从表格单元格中提取文本内容
    """
    content_parts = []
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            content_parts.append(paragraph.text.strip())
    return " ".join(content_parts)


def extract_docx_content_simple(file_path):
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return paragraphs
    except Exception as e:
        logger.error(f"读取文档失败: {e}")
        return []


def ignored_sentence(sentence, language='en'):
    """
    忽略的句子，不进行对齐
    """
    if language == 'en' and len(sentence) <= EN_SENTENCE_IGNORE_LENGTH:
        return True
    if language == 'zh' and len(sentence) <= ZH_SENTENCE_IGNORE_LENGTH:
        return True
    if len(sentence.splitlines()) > 1:
        return True
    if re.match(digit_only_pattern, sentence):
        return True
    if language == 'zh' and re.match(zh_digit_english_only_pattern, sentence):
        return True

    return False


def extract_docx_content(file_path:str, file_language:str='en'):
    """
    提取 .docx 文件中的所有文本内容
    """
    try:
        document = docx.Document(file_path)

        all_content = []
        # 提取段落内容
        for paragraph in document.paragraphs:
            temp_content = paragraph.text.strip()
            if temp_content \
                    and not ignored_sentence(temp_content, file_language) \
                    and temp_content not in all_content:
                # logger.info(temp_content)
                all_content.append(temp_content)
        # 提取表格内容
        for element in document.element.body:
            if isinstance(element, CT_Tbl):
                table = Table(element, document)
                for row in table.rows:
                    for cell in row.cells:
                        cell_content = extract_paragraphs_from_cell(cell)
                        temp_content = cell_content.strip()
                        if temp_content \
                                and not ignored_sentence(temp_content, file_language) \
                                and temp_content not in all_content:
                            # logger.info(temp_content)
                            all_content.append(temp_content)

        return all_content

    except Exception as e:
        logger.error(f"读取文档失败: {e}")
        return []
